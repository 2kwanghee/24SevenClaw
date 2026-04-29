#!/usr/bin/env python3
"""24SevenClaw 프로젝트 사양 문서(.docx) 자동 생성 스크립트.

docs/spec/ 디렉토리에 기획서, API 정의서, 기술설계서, 시스템 아키텍처 문서를 생성한다.
매일 실행하여 최신 코드/문서 상태를 반영한 문서를 갱신할 수 있다.

사용법:
    python scripts/generate_spec_docs.py [--all | --prd | --api | --tech | --arch | --daily]
"""

from __future__ import annotations

import argparse
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

# --- 경로 설정 ---
ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
SPEC_DIR = DOCS_DIR / "spec"
SPEC_DIR.mkdir(parents=True, exist_ok=True)

TODAY = datetime.now().strftime("%Y-%m-%d")
VERSION = "0.1.0"


# ============================================================
# 유틸리티
# ============================================================

def _style_doc(doc: Document) -> None:
    """문서 공통 스타일 설정."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "맑은 고딕"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_cover(doc: Document, title: str, subtitle: str) -> None:
    """표지 페이지 추가."""
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("24SevenClaw")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(20)
    run2.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"버전 {VERSION}  |  {TODAY}")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """표 추가."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")


def _add_code_block(doc: Document, code: str) -> None:
    """코드 블록 추가 (고정폭 폰트)."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(1)


# ============================================================
# 1. 기획서 (PRD)
# ============================================================

def generate_prd() -> Path:
    """기획서 (Product Requirements Document) 생성."""
    doc = Document()
    _style_doc(doc)
    _add_cover(doc, "기획서 (PRD)", "라이센스 기반 AI 에이전트 개발 오케스트레이션 플랫폼")

    # 목차
    doc.add_heading("목차", level=1)
    toc_items = [
        "1. 프로젝트 개요",
        "2. 비즈니스 모델",
        "3. 타깃 사용자",
        "4. 핵심 기능 요구사항",
        "5. 레포지토리 구조",
        "6. 기술 스택",
        "7. 개발 로드맵",
        "8. 경쟁 분석",
        "9. 라이센스 모델",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. 프로젝트 개요
    doc.add_heading("1. 프로젝트 개요", level=1)
    doc.add_paragraph(
        "24SevenClaw는 라이센스 기반 AI 에이전트 개발 오케스트레이션 플랫폼이다. "
        "클라우드(컨트롤 플레인)에서 프로젝트를 관리하고, 고객 서버(실행 플레인)에서 "
        "실제 코드 실행과 데이터 저장을 수행하는 분리 아키텍처를 채택한다."
    )
    doc.add_paragraph(
        "이 모델은 GitHub Actions Self-hosted Runner, Ansible Tower/AWX, "
        "GitLab Runner와 동일한 패턴을 따르며, 고객의 데이터 주권을 보장한다."
    )

    doc.add_heading("핵심 가치", level=2)
    values = [
        "클라우드 UI로 고객 서버의 개발 환경을 원격 구성",
        "개발 프로세스 오케스트레이션 (티켓 → 코드 → 빌드 → 배포)",
        "고객 코드/데이터는 절대 클라우드로 전송하지 않음 (데이터 주권)",
        "에이전트/스킬/MCP 레지스트리를 통한 확장 가능한 생태계",
    ]
    for v in values:
        doc.add_paragraph(v, style="List Bullet")

    # 2. 비즈니스 모델
    doc.add_heading("2. 비즈니스 모델", level=1)
    _add_table(doc, ["구분", "설명"], [
        ["우리 (24SevenClaw)", "클라우드 웹서비스(SaaS) + 라이센스 제공"],
        ["고객", "자사 서버에서 실행, 모든 데이터는 고객 소유"],
        ["라이센스", "프로젝트 단위 (에이전트/스킬 포함)"],
        ["수익원", "구독 기반 라이센스 (Free / Pro / Enterprise)"],
    ])

    # 3. 타깃 사용자
    doc.add_heading("3. 타깃 사용자", level=1)
    doc.add_paragraph(
        "AI 기반 개발 자동화를 도입하고자 하는 개발팀 및 기업. "
        "특히 보안/컴플라이언스 요구사항으로 인해 코드를 외부 클라우드에 "
        "올릴 수 없는 조직을 주요 타깃으로 한다."
    )

    # 4. 핵심 기능 요구사항
    doc.add_heading("4. 핵심 기능 요구사항", level=1)

    doc.add_heading("4.1 컨트롤 플레인 (Cloud)", level=2)
    cloud_features = [
        ["사용자 인증", "계정 관리, JWT 기반 로그인/회원가입"],
        ["라이센스 관리", "프로젝트 단위 라이센스 발급/검증/갱신"],
        ["레지스트리", "에이전트/스킬/MCP 서버 카탈로그 관리"],
        ["프로젝트 관리", "프로젝트 CRUD, 설정 매핑, 환경 구성"],
        ["티켓/이슈", "개발 작업 발행, 트래킹, 진행 모니터링"],
        ["오케스트레이션", "Agent에게 명령 전달, 상태 수집"],
        ["모니터링", "Agent 연결 상태, 작업 진행 현황 대시보드"],
    ]
    _add_table(doc, ["기능", "설명"], cloud_features)

    doc.add_heading("4.2 실행 플레인 (Agent)", level=2)
    agent_features = [
        ["클라우드 연결", "WebSocket(WSS)으로 컨트롤 플레인에 아웃바운드 연결"],
        ["Docker 관리", "컨테이너 생성/시작/중지/삭제 라이프사이클"],
        ["환경 프로비저닝", "설정 기반 개발 환경 자동 구성 (템플릿 시스템)"],
        ["Claude 관리", "Claude 인스턴스 실행, 작업 전달, 결과 수집"],
        ["Git 관리", "저장소 초기화, 커밋, 푸시"],
        ["빌드/실행", "코드 빌드, 서버 실행, 로그 수집"],
        ["상태 보고", "진행 상황을 클라우드로 실시간 보고"],
    ]
    _add_table(doc, ["기능", "설명"], agent_features)

    # 5. 레포지토리 구조
    doc.add_heading("5. 레포지토리 구조", level=1)
    repos = [
        ["clickeye-web", "Next.js 15", "3000", "클라우드 프론트엔드"],
        ["clickeye-api", "FastAPI", "8000", "클라우드 백엔드"],
        ["clickeye-agent", "Python", "-", "고객 서버 에이전트 데몬"],
        ["clickeye-infra", "Docker/YAML", "-", "인프라 설정"],
        ["clickeye-contracts", "TypeScript", "-", "공유 타입/프로토콜"],
    ]
    _add_table(doc, ["레포", "기술", "포트", "역할"], repos)

    # 6. 기술 스택
    doc.add_heading("6. 기술 스택", level=1)
    stack = [
        ["Frontend", "Next.js 15 (App Router), TypeScript, Tailwind CSS, shadcn/ui, Zustand, TanStack Query v5"],
        ["Backend (Cloud)", "FastAPI 0.115+, Python 3.12+, SQLAlchemy 2.0 async, Alembic, Pydantic v2"],
        ["Backend (Agent)", "Python 3.12+, websockets, docker-py, asyncio, aiosqlite"],
        ["Database (Cloud)", "PostgreSQL 16, Redis 7"],
        ["Database (Agent)", "SQLite (로컬 상태 관리)"],
        ["통신", "WebSocket (WSS/TLS), Agent → Cloud 아웃바운드"],
        ["인증", "Auth.js v5 (프론트), JWT python-jose (백엔드)"],
        ["CI/CD", "GitHub Actions"],
        ["패키지 관리", "uv (Python), npm (Node)"],
    ]
    _add_table(doc, ["영역", "기술"], stack)

    # 7. 개발 로드맵
    doc.add_heading("7. 개발 로드맵", level=1)
    doc.add_paragraph("시작일: 2026-03-23 | 총 20주 (5 Phase)")

    phases = [
        ["Phase 0", "Week 1-2", "프로젝트 셋업", "레포 초기화, 인증, 기본 CRUD, CI/CD, DB"],
        ["Phase 1", "Week 3-5", "MVP Core + Agent 통신", "레지스트리, WebSocket 프로토콜, Agent 등록"],
        ["Phase 2", "Week 6-9", "Docker 프로비저닝", "컨테이너 관리, 환경 템플릿, 모니터링"],
        ["Phase 3", "Week 10-13", "티켓 시스템 + Claude", "티켓 CRUD, Claude 연동, 실시간 진행"],
        ["Phase 4", "Week 14-16", "파이프라인 자동화", "DAG, 빌더 UI, 트리거, 빌드/실행"],
        ["Phase 5", "Week 17-20", "라이센스 + 상용화", "Stripe 결제, 관리자, 랜딩, 보안 감사"],
    ]
    _add_table(doc, ["Phase", "기간", "이름", "주요 항목"], phases)

    # 8. 경쟁 분석
    doc.add_heading("8. 경쟁 분석", level=1)
    comp = [
        ["GitHub Actions Runner", "CI/CD 자동화", "워크플로 중심, AI 에이전트 미지원"],
        ["Ansible Tower/AWX", "인프라 자동화", "서버 구성 관리 중심, 개발 프로세스 미지원"],
        ["GitLab Runner", "CI/CD 실행기", "빌드/배포 중심, AI 개발 오케스트레이션 미지원"],
    ]
    _add_table(doc, ["플랫폼", "핵심 역할", "24SevenClaw와의 차이"], comp)

    doc.add_paragraph(
        "24SevenClaw의 차별점: AI 에이전트 중심 개발 오케스트레이션 + "
        "에이전트/스킬/MCP 레지스트리 생태계 + 고객 데이터 주권 보장."
    )

    # 9. 라이센스 모델
    doc.add_heading("9. 라이센스 모델", level=1)
    license_table = [
        ["프로젝트 수", "1", "5", "Unlimited"],
        ["Agent 연결", "1", "3", "Unlimited"],
        ["에이전트/스킬/MCP", "기본 세트", "전체", "전체 + 커스텀"],
        ["동시 티켓", "1", "10", "Unlimited"],
        ["환경 템플릿", "기본", "전체", "전체 + 커스텀"],
        ["기술 지원", "커뮤니티", "이메일", "전담"],
    ]
    _add_table(doc, ["항목", "Free", "Pro", "Enterprise"], license_table)

    out = SPEC_DIR / f"24SevenClaw_기획서_PRD_{TODAY}.docx"
    doc.save(str(out))
    return out


# ============================================================
# 2. API 정의서
# ============================================================

def generate_api_spec() -> Path:
    """API 정의서 생성."""
    doc = Document()
    _style_doc(doc)
    _add_cover(doc, "API 정의서", "REST API + WebSocket 프로토콜 사양")

    # 목차
    doc.add_heading("목차", level=1)
    for item in [
        "1. API 개요",
        "2. 인증 (Authentication)",
        "3. REST API 엔드포인트",
        "4. WebSocket 프로토콜",
        "5. 에러 처리",
        "6. 데이터 모델",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. API 개요
    doc.add_heading("1. API 개요", level=1)
    doc.add_paragraph("Base URL: https://api.24sevenclaw.com/api/v1")
    doc.add_paragraph("WebSocket URL: wss://api.24sevenclaw.com/ws/agent")
    doc.add_paragraph("Content-Type: application/json")
    doc.add_paragraph("인증: Bearer JWT Token (REST), agent_secret (WebSocket)")

    # 2. 인증
    doc.add_heading("2. 인증 (Authentication)", level=1)

    doc.add_heading("2.1 사용자 인증 (JWT)", level=2)
    doc.add_paragraph(
        "REST API는 JWT 기반 인증을 사용한다. "
        "access_token (15분) + refresh_token (7일) 방식."
    )

    doc.add_heading("2.2 Agent 인증", level=2)
    doc.add_paragraph(
        "Agent는 최초 등록 시 registration_token을 사용하여 agent_secret을 발급받는다. "
        "이후 WebSocket 연결 시 Authorization 헤더에 agent_secret을 Bearer 토큰으로 전달한다."
    )

    # 3. REST API 엔드포인트
    doc.add_heading("3. REST API 엔드포인트", level=1)

    doc.add_heading("3.1 Health Check", level=2)
    _add_table(doc, ["항목", "값"], [
        ["Method", "GET"],
        ["Path", "/health"],
        ["인증", "불필요"],
        ["설명", "서버 및 DB 연결 상태 확인"],
    ])
    doc.add_paragraph("응답 예시:")
    _add_code_block(doc, '{"status": "healthy", "version": "0.1.0", "db": "connected"}')

    doc.add_heading("3.2 회원가입", level=2)
    _add_table(doc, ["항목", "값"], [
        ["Method", "POST"],
        ["Path", "/api/v1/auth/register"],
        ["인증", "불필요"],
        ["설명", "새 사용자 계정 생성"],
    ])
    doc.add_paragraph("요청 본문:")
    _add_code_block(doc, '{\n  "email": "user@example.com",\n  "password": "securepass123",\n  "display_name": "홍길동"\n}')
    doc.add_paragraph("응답 (201 Created):")
    _add_code_block(doc, '{\n  "id": "uuid",\n  "email": "user@example.com",\n  "display_name": "홍길동",\n  "avatar_url": null,\n  "plan": "free",\n  "created_at": "2026-03-23T10:00:00Z"\n}')

    doc.add_heading("3.3 로그인", level=2)
    _add_table(doc, ["항목", "값"], [
        ["Method", "POST"],
        ["Path", "/api/v1/auth/login"],
        ["인증", "불필요"],
        ["설명", "이메일/비밀번호로 JWT 토큰 발급"],
    ])
    doc.add_paragraph("요청 본문:")
    _add_code_block(doc, '{\n  "email": "user@example.com",\n  "password": "securepass123"\n}')
    doc.add_paragraph("응답:")
    _add_code_block(doc, '{\n  "access_token": "eyJ...",\n  "refresh_token": "eyJ...",\n  "token_type": "bearer"\n}')

    doc.add_heading("3.4 토큰 갱신", level=2)
    _add_table(doc, ["항목", "값"], [
        ["Method", "POST"],
        ["Path", "/api/v1/auth/refresh"],
        ["인증", "불필요"],
        ["설명", "refresh_token으로 새 토큰 쌍 발급"],
    ])

    doc.add_heading("3.5 현재 사용자 조회", level=2)
    _add_table(doc, ["항목", "값"], [
        ["Method", "GET"],
        ["Path", "/api/v1/auth/me"],
        ["인증", "Bearer Token"],
        ["설명", "현재 인증된 사용자 정보 조회"],
    ])

    doc.add_heading("3.6 향후 구현 예정 API", level=2)
    planned = [
        ["Projects CRUD", "/api/v1/projects", "프로젝트 목록/생성/수정/삭제", "Phase 0"],
        ["Registry - Agents", "/api/v1/registry/agents", "에이전트 카탈로그 조회", "Phase 1"],
        ["Registry - Skills", "/api/v1/registry/skills", "스킬 카탈로그 조회", "Phase 1"],
        ["Registry - MCPs", "/api/v1/registry/mcps", "MCP 서버 카탈로그 조회", "Phase 1"],
        ["Project Config", "/api/v1/projects/{id}/config", "프로젝트별 설정 관리", "Phase 1"],
        ["Tickets", "/api/v1/tickets", "개발 티켓 CRUD", "Phase 3"],
        ["Agent Status", "/api/v1/agents", "Agent 연결 상태 조회", "Phase 1"],
    ]
    _add_table(doc, ["기능", "경로", "설명", "Phase"], planned)

    # 4. WebSocket 프로토콜
    doc.add_heading("4. WebSocket 프로토콜", level=1)

    doc.add_heading("4.1 연결 정보", level=2)
    _add_table(doc, ["항목", "값"], [
        ["URL", "wss://api.24sevenclaw.com/ws/agent?agent_id={agent_id}"],
        ["인증", "Authorization: Bearer {agent_secret}"],
        ["Heartbeat", "30초 간격"],
        ["재연결", "지수 백오프 (1s → 2s → 4s → ... → max 300s)"],
    ])

    doc.add_heading("4.2 메시지 Envelope", level=2)
    _add_code_block(doc, (
        '{\n'
        '  "id": "msg_uuid_v4",\n'
        '  "type": "agent.heartbeat | command.setup_env | ...",\n'
        '  "timestamp": "2026-03-23T10:30:00Z",\n'
        '  "payload": { ... },\n'
        '  "signature": "hmac_sha256_hex"  // agent_secret으로 서명\n'
        '}'
    ))

    doc.add_heading("4.3 Agent → Cloud 메시지", level=2)
    a2c_msgs = [
        ["agent.register", "최초 등록 요청", "registration_token, hostname, os, docker_version, capabilities"],
        ["agent.heartbeat", "상태 보고 (30초)", "status, uptime, system(cpu/mem/disk), environments, active_tasks"],
        ["agent.status", "이벤트 알림", "event, project_id, task_id, detail"],
        ["agent.log", "로그 스트리밍", "project_id, task_id, level, source, message"],
        ["agent.result", "작업 완료 결과", "task_id, ticket_id, status, summary, changes, metrics"],
    ]
    _add_table(doc, ["메시지 타입", "설명", "주요 페이로드"], a2c_msgs)

    doc.add_heading("4.4 Cloud → Agent 메시지", level=2)
    c2a_msgs = [
        ["command.setup_env", "환경 프로비저닝", "project_id, environment(template, agents, skills, mcps, claude), git"],
        ["command.deploy_ticket", "티켓 전달", "ticket_id, project_id, title, description, acceptance_criteria"],
        ["command.build", "빌드 실행", "project_id, build_type, command, env_vars"],
        ["command.run", "서비스 실행", "project_id, command, port, env_vars"],
        ["command.stop", "서비스 중지", "project_id, target, force"],
        ["command.destroy_env", "환경 삭제", "project_id, keep_git, keep_data"],
        ["config.update", "설정 변경", "project_id, changes[{action, target, id, config}]"],
    ]
    _add_table(doc, ["메시지 타입", "설명", "주요 페이로드"], c2a_msgs)

    doc.add_heading("4.5 서명 검증", level=2)
    doc.add_paragraph("모든 메시지는 HMAC-SHA256으로 서명된다.")
    _add_code_block(doc, "signature = HMAC-SHA256(key=agent_secret, message=id+type+timestamp+JSON(payload))")

    # 5. 에러 처리
    doc.add_heading("5. 에러 처리", level=1)
    errors = [
        ["AUTH_FAILED", "인증 실패", "No"],
        ["LICENSE_EXPIRED", "라이센스 만료", "No"],
        ["LICENSE_LIMIT", "라이센스 한도 초과", "No"],
        ["DOCKER_ERROR", "Docker 작업 실패", "Yes"],
        ["ENV_SETUP_FAILED", "환경 구성 실패", "Yes"],
        ["GIT_ERROR", "Git 작업 실패", "Yes"],
        ["CLAUDE_ERROR", "Claude 작업 실패", "Yes"],
        ["BUILD_FAILED", "빌드 실패", "Yes"],
        ["RESOURCE_LIMIT", "서버 리소스 부족", "No"],
        ["TIMEOUT", "작업 시간 초과", "Yes"],
    ]
    _add_table(doc, ["에러 코드", "설명", "복구 가능"], errors)

    # 6. 데이터 모델
    doc.add_heading("6. 데이터 모델", level=1)

    doc.add_heading("6.1 UserCreate (요청)", level=2)
    _add_table(doc, ["필드", "타입", "필수", "설명"], [
        ["email", "EmailStr", "O", "이메일 주소"],
        ["password", "str", "O", "비밀번호 (최소 8자)"],
        ["display_name", "str", "O", "표시 이름 (1~100자)"],
    ])

    doc.add_heading("6.2 UserResponse (응답)", level=2)
    _add_table(doc, ["필드", "타입", "설명"], [
        ["id", "UUID", "사용자 고유 ID"],
        ["email", "str", "이메일 주소"],
        ["display_name", "str", "표시 이름"],
        ["avatar_url", "str | null", "아바타 URL"],
        ["plan", "str", "라이센스 플랜 (free/pro/enterprise)"],
        ["created_at", "datetime", "생성 일시"],
    ])

    doc.add_heading("6.3 TokenResponse (응답)", level=2)
    _add_table(doc, ["필드", "타입", "설명"], [
        ["access_token", "str", "JWT 액세스 토큰 (~15분)"],
        ["refresh_token", "str", "JWT 리프레시 토큰 (~7일)"],
        ["token_type", "str", "bearer"],
    ])

    out = SPEC_DIR / f"24SevenClaw_API정의서_{TODAY}.docx"
    doc.save(str(out))
    return out


# ============================================================
# 3. 기술설계서
# ============================================================

def generate_tech_design() -> Path:
    """기술설계서 생성."""
    doc = Document()
    _style_doc(doc)
    _add_cover(doc, "기술설계서", "시스템 구현 설계 및 기술 상세")

    doc.add_heading("목차", level=1)
    for item in [
        "1. 설계 원칙",
        "2. 백엔드 설계 (FastAPI)",
        "3. 프론트엔드 설계 (Next.js)",
        "4. Agent 데몬 설계",
        "5. 데이터베이스 설계",
        "6. 보안 설계",
        "7. 통신 설계",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. 설계 원칙
    doc.add_heading("1. 설계 원칙", level=1)
    principles = [
        ["비동기 우선", "모든 I/O 작업은 async/await 사용 (FastAPI, Agent)"],
        ["서비스 레이어 분리", "엔드포인트(얇음) → 서비스(비즈니스 로직) → 모델(데이터)"],
        ["타입 안전성", "Python: strict mypy, TypeScript: strict mode"],
        ["Contract 우선", "API 변경 시 contracts 레포 스키마를 먼저 업데이트"],
        ["데이터 주권", "고객 코드/데이터는 절대 클라우드에 전송하지 않음"],
    ]
    _add_table(doc, ["원칙", "설명"], principles)

    # 2. 백엔드 설계
    doc.add_heading("2. 백엔드 설계 (FastAPI)", level=1)

    doc.add_heading("2.1 디렉토리 구조", level=2)
    _add_code_block(doc, (
        "app/\n"
        "├── main.py              # FastAPI app factory\n"
        "├── config.py            # Pydantic BaseSettings\n"
        "├── database.py          # async SQLAlchemy engine + session\n"
        "├── dependencies.py      # DI (get_db, get_current_user)\n"
        "├── api/v1/              # API 엔드포인트\n"
        "│   ├── router.py        # v1 라우터 집합\n"
        "│   ├── auth.py          # 인증 (register, login, refresh, me)\n"
        "│   └── health.py        # 헬스 체크\n"
        "├── models/              # SQLAlchemy ORM 모델\n"
        "├── schemas/             # Pydantic 요청/응답 스키마\n"
        "├── services/            # 비즈니스 로직\n"
        "├── core/                # JWT, 예외, 미들웨어\n"
        "├── utils/               # 페이지네이션, 캐시\n"
        "└── ws/                  # WebSocket Hub (Agent 통신)\n"
    ))

    doc.add_heading("2.2 인증 흐름", level=2)
    doc.add_paragraph(
        "JWT 기반 인증: access_token (~15분) + refresh_token (~7일). "
        "비밀번호는 bcrypt 해싱. python-jose로 JWT 생성/검증."
    )

    doc.add_heading("2.3 WebSocket Hub 설계", level=2)
    doc.add_paragraph(
        "Agent 연결 관리를 위한 WebSocket Hub. "
        "Redis Pub/Sub로 다중 API 인스턴스 간 메시지 브로커링. "
        "메시지별 HMAC-SHA256 서명 검증."
    )

    # 3. 프론트엔드 설계
    doc.add_heading("3. 프론트엔드 설계 (Next.js)", level=1)

    doc.add_heading("3.1 기술 스택", level=2)
    _add_table(doc, ["기술", "용도"], [
        ["Next.js 15 (App Router)", "프레임워크, SSR/SSG"],
        ["React 19", "UI 라이브러리"],
        ["TypeScript (strict)", "타입 안전성"],
        ["Tailwind CSS", "유틸리티 CSS"],
        ["shadcn/ui", "컴포넌트 라이브러리 (커스터마이즈 가능)"],
        ["Zustand", "클라이언트 상태 관리 (UI 상태만)"],
        ["TanStack Query v5", "서버 상태 관리 (API 호출/캐싱)"],
        ["Auth.js v5", "인증 (JWT)"],
    ])

    doc.add_heading("3.2 주요 페이지", level=2)
    pages = [
        ["/", "랜딩 페이지"],
        ["/login, /register", "인증 페이지"],
        ["/projects", "프로젝트 목록"],
        ["/projects/new", "프로젝트 생성"],
        ["/projects/[id]/configure", "프로젝트 설정 (에이전트/스킬/MCP)"],
        ["/projects/[id]/settings", "프로젝트 기본 설정"],
        ["/registry/agents", "에이전트 레지스트리 브라우저"],
        ["/registry/skills", "스킬 레지스트리 브라우저"],
        ["/registry/mcps", "MCP 서버 레지스트리 브라우저"],
    ]
    _add_table(doc, ["경로", "설명"], pages)

    # 4. Agent 데몬 설계
    doc.add_heading("4. Agent 데몬 설계", level=1)

    doc.add_heading("4.1 핸들러 아키텍처", level=2)
    handlers = [
        ["DockerHandler", "Docker 컨테이너 생성/시작/중지/삭제"],
        ["EnvHandler", "환경 프로비저닝 (템플릿 → Docker Compose)"],
        ["ClaudeHandler", "Claude 인스턴스 관리, 작업 전달 (Phase 3)"],
        ["GitHandler", "Git 저장소 초기화, 커밋, 푸시 (Phase 2)"],
        ["BuildHandler", "빌드 실행, 로그 수집 (Phase 4)"],
        ["PipelineHandler", "파이프라인 DAG 실행 (Phase 4)"],
    ]
    _add_table(doc, ["핸들러", "책임"], handlers)

    doc.add_heading("4.2 연결 관리", level=2)
    doc.add_paragraph(
        "CloudConnection 클래스가 WebSocket 연결을 관리한다. "
        "Dispatcher가 메시지 타입별로 적절한 핸들러에 라우팅. "
        "Reporter가 상태를 주기적으로 Cloud에 보고."
    )

    doc.add_heading("4.3 로컬 상태 저장", level=2)
    doc.add_paragraph(
        "SQLite(aiosqlite)를 사용하여 Agent 로컬 상태를 관리한다. "
        "환경 목록, 작업 히스토리, 연결 설정 등을 저장."
    )

    # 5. 데이터베이스 설계
    doc.add_heading("5. 데이터베이스 설계", level=1)

    doc.add_heading("5.1 Cloud DB (PostgreSQL 16)", level=2)
    db_tables = [
        ["users", "사용자 계정", "id, email, password_hash, display_name, avatar_url, plan, is_active"],
        ["licenses", "프로젝트 라이센스", "id, user_id, license_key, plan, limits, valid_from/until, status"],
        ["projects", "프로젝트 메타데이터", "id, owner_id, name, slug, settings"],
        ["agents", "에이전트 레지스트리", "id, name, image, config, description"],
        ["skills", "스킬 레지스트리", "id, name, image, config, description"],
        ["mcp_servers", "MCP 서버 레지스트리", "id, name, image, config, description"],
        ["project_configs", "프로젝트별 설정", "id, project_id, agents[], skills[], mcps[]"],
        ["agent_connections", "Agent 연결 상태", "id, agent_id, license_id, secret_hash, hostname, status"],
        ["tickets", "개발 티켓", "id, project_id, title, description, priority, status"],
        ["ticket_events", "티켓 이벤트 로그", "id, ticket_id, event_type, data, created_at"],
    ]
    _add_table(doc, ["테이블", "설명", "주요 컬럼"], db_tables)

    doc.add_heading("5.2 마이그레이션", level=2)
    doc.add_paragraph(
        "Alembic (async) 사용. 모든 스키마 변경은 마이그레이션 파일을 통해 관리. "
        "수동 DB 스키마 변경 금지."
    )

    # 6. 보안 설계
    doc.add_heading("6. 보안 설계", level=1)
    security = [
        ["통신 암호화", "모든 WebSocket/HTTP 통신은 TLS (WSS/HTTPS)"],
        ["메시지 서명", "HMAC-SHA256으로 모든 WebSocket 메시지 서명"],
        ["비밀번호 해싱", "bcrypt (passlib)"],
        ["JWT 토큰", "access_token (~15분) + refresh_token (~7일)"],
        ["Agent 인증", "registration_token (1회) → agent_secret (영구, 로컬 저장)"],
        ["데이터 분리", "고객 코드/데이터는 클라우드에 전송하지 않음"],
        ["민감 정보", "API 키 등은 환경 변수 참조(_env 접미사)로 처리"],
    ]
    _add_table(doc, ["항목", "설명"], security)

    # 7. 통신 설계
    doc.add_heading("7. 통신 설계", level=1)
    doc.add_paragraph(
        "Agent → Cloud 방향의 아웃바운드 WebSocket 연결. "
        "고객 서버의 방화벽 뒤에서도 동작 (인바운드 포트 오픈 불필요). "
        "HTTPS/WSS(443) 포트만 사용."
    )

    doc.add_heading("7.1 연결 라이프사이클", level=2)
    lifecycle = [
        ["1", "최초 등록", "Cloud UI에서 agent_id + registration_token 발급"],
        ["2", "Agent 설치", "install.sh --token <registration_token>"],
        ["3", "WebSocket 연결", "Agent → Cloud 아웃바운드 연결"],
        ["4", "등록 메시지", "agent.register 전송 → agent_secret 수신"],
        ["5", "Heartbeat", "30초 간격 상태 보고"],
        ["6", "재연결", "끊김 시 지수 백오프 (1s → max 300s)"],
    ]
    _add_table(doc, ["순서", "단계", "설명"], lifecycle)

    out = SPEC_DIR / f"24SevenClaw_기술설계서_{TODAY}.docx"
    doc.save(str(out))
    return out


# ============================================================
# 4. 시스템 아키텍처
# ============================================================

def generate_architecture() -> Path:
    """시스템 아키텍처 문서 생성."""
    doc = Document()
    _style_doc(doc)
    _add_cover(doc, "시스템 아키텍처", "컨트롤 플레인 / 실행 플레인 분리 아키텍처")

    doc.add_heading("목차", level=1)
    for item in [
        "1. 아키텍처 개요",
        "2. 컨트롤 플레인 (Cloud)",
        "3. 실행 플레인 (Customer Server)",
        "4. 통신 아키텍처",
        "5. 데이터 흐름",
        "6. 보안 아키텍처",
        "7. 확장성 설계",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. 아키텍처 개요
    doc.add_heading("1. 아키텍처 개요", level=1)
    doc.add_paragraph(
        "24SevenClaw는 컨트롤 플레인(Cloud)과 실행 플레인(Customer Server)이 "
        "분리된 아키텍처를 채택한다. 클라우드에서 프로젝트 관리/오케스트레이션을 수행하고, "
        "고객 서버에서 실제 코드 실행과 데이터 저장을 수행한다."
    )

    doc.add_heading("전체 아키텍처 다이어그램", level=2)
    _add_code_block(doc, (
        "┌──────────────────────────────────────────┐\n"
        "│        24SevenClaw Cloud                  │\n"
        "│        (컨트롤 플레인)                      │\n"
        "│                                          │\n"
        "│  [Next.js Frontend] ←→ [FastAPI Backend] │\n"
        "│                           │              │\n"
        "│              [PostgreSQL] + [Redis]       │\n"
        "└───────────────────┬──────────────────────┘\n"
        "                    │ WebSocket (TLS)\n"
        "                    │ Agent → Cloud\n"
        "        ┌───────────┼───────────┐\n"
        "        ▼           ▼           ▼\n"
        "  [고객사 A]   [고객사 B]   [고객사 C]\n"
        "  ┌────────┐  ┌────────┐  ┌────────┐\n"
        "  │ Agent  │  │ Agent  │  │ Agent  │\n"
        "  │ 데몬   │  │ 데몬   │  │ 데몬   │\n"
        "  │   ↓    │  │   ↓    │  │   ↓    │\n"
        "  │ Docker │  │ Docker │  │ Docker │\n"
        "  │ Engine │  │ Engine │  │ Engine │\n"
        "  └────────┘  └────────┘  └────────┘\n"
    ))

    # 2. 컨트롤 플레인
    doc.add_heading("2. 컨트롤 플레인 (Cloud)", level=1)

    doc.add_heading("2.1 역할과 책임", level=2)
    roles = [
        ["사용자 인증", "계정 관리, 로그인, JWT 토큰"],
        ["라이센스 관리", "프로젝트 단위 라이센스 발급/검증"],
        ["레지스트리", "에이전트/스킬/MCP 카탈로그 (회사 IP)"],
        ["프로젝트 관리", "프로젝트 메타데이터, 설정 참조"],
        ["티켓/이슈", "개발 작업 발행, 트래킹"],
        ["오케스트레이션", "Agent에게 명령 전달, 상태 수집"],
        ["모니터링", "Agent 연결 상태, 작업 진행 현황"],
    ]
    _add_table(doc, ["역할", "설명"], roles)

    doc.add_heading("2.2 저장 데이터 정책", level=2)
    doc.add_heading("저장하는 데이터", level=3)
    stores = [
        "사용자 계정 (email, password_hash)",
        "라이센스 정보 (프로젝트별)",
        "프로젝트 메타데이터 (이름, 설명, 설정 참조)",
        "에이전트/스킬/MCP 레지스트리 (우리 IP)",
        "프로젝트별 에이전트/스킬/MCP 설정 매핑",
        "Agent 연결 정보 (ID, 마지막 접속, 상태)",
        "티켓/이슈 (제목, 설명, 상태, 우선순위)",
    ]
    for s in stores:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("저장하지 않는 데이터", level=3)
    not_stores = [
        "고객 소스 코드",
        "빌드 결과물/아티팩트",
        "실행 로그 원본 (요약만 수신)",
        "Git 저장소 데이터",
        "고객 환경 설정 파일 원본",
        "고객의 비즈니스 데이터",
    ]
    for s in not_stores:
        doc.add_paragraph(s, style="List Bullet")

    # 3. 실행 플레인
    doc.add_heading("3. 실행 플레인 (Customer Server)", level=1)

    doc.add_heading("3.1 Agent 데몬 구조", level=2)
    _add_code_block(doc, (
        "clickeye-agent/\n"
        "├── agent/\n"
        "│   ├── main.py            # 데몬 엔트리포인트\n"
        "│   ├── config.py          # 설정 (라이센스 키, 클라우드 URL)\n"
        "│   ├── connection.py      # WebSocket 클라이언트\n"
        "│   ├── dispatcher.py      # 메시지 → 핸들러 라우팅\n"
        "│   ├── handlers/\n"
        "│   │   ├── docker_handler.py\n"
        "│   │   ├── env_handler.py\n"
        "│   │   ├── claude_handler.py\n"
        "│   │   ├── git_handler.py\n"
        "│   │   ├── build_handler.py\n"
        "│   │   └── pipeline_handler.py\n"
        "│   ├── reporter.py        # 상태 보고\n"
        "│   └── local_store.py     # SQLite 로컬 상태\n"
        "└── templates/             # Docker Compose 환경 템플릿\n"
    ))

    # 4. 통신 아키텍처
    doc.add_heading("4. 통신 아키텍처", level=1)
    doc.add_paragraph(
        "Agent가 Cloud로 아웃바운드 WebSocket 연결을 맺는다. "
        "고객 서버의 방화벽 뒤에서도 동작하며 인바운드 포트 오픈이 불필요하다."
    )

    doc.add_heading("4.1 프로토콜 요약", level=2)

    doc.add_paragraph("Agent → Cloud:")
    a2c = ["agent.register (최초 등록)", "agent.heartbeat (30초 상태 보고)",
           "agent.status (이벤트 알림)", "agent.log (로그 요약 스트리밍)",
           "agent.result (작업 완료 결과)"]
    for m in a2c:
        doc.add_paragraph(m, style="List Bullet")

    doc.add_paragraph("Cloud → Agent:")
    c2a = ["command.setup_env (환경 프로비저닝)", "command.deploy_ticket (티켓 전달)",
           "command.build / run / stop (빌드/실행 제어)",
           "command.destroy_env (환경 삭제)", "config.update (설정 변경)"]
    for m in c2a:
        doc.add_paragraph(m, style="List Bullet")

    # 5. 데이터 흐름
    doc.add_heading("5. 데이터 흐름", level=1)

    doc.add_heading("5.1 환경 셋업 흐름", level=2)
    _add_code_block(doc, (
        "[Cloud UI: 환경 셋업 명령]\n"
        "        ↓\n"
        "[Agent: WebSocket으로 수신]\n"
        "        ↓\n"
        "[docker_handler: 컨테이너 생성]\n"
        "  ├── 에이전트 런타임 이미지 pull & run\n"
        "  ├── 스킬 서버 이미지 pull & run\n"
        "  ├── MCP 서버 이미지 pull & run\n"
        "  └── Claude 설치 & 설정\n"
        "        ↓\n"
        "[git_handler: Git 저장소 초기화]\n"
        "        ↓\n"
        "[reporter: 완료 상태를 Cloud로 보고]\n"
    ))

    doc.add_heading("5.2 티켓 처리 흐름", level=2)
    _add_code_block(doc, (
        "[Cloud UI: 개발 티켓 발행]\n"
        "        ↓\n"
        "[Agent: WebSocket으로 수신]\n"
        "        ↓\n"
        "[claude_handler: Claude에 작업 지시]\n"
        "        ↓\n"
        "[Claude: 코드 작성/수정]\n"
        "        ↓\n"
        "[git_handler: 변경사항 커밋 → 고객 Git 푸시]\n"
        "        ↓\n"
        "[reporter: 작업 결과 요약을 Cloud로 보고]\n"
    ))

    # 6. 보안 아키텍처
    doc.add_heading("6. 보안 아키텍처", level=1)
    sec = [
        ["통신 보안", "TLS (WSS/HTTPS), 모든 메시지 HMAC-SHA256 서명"],
        ["인증", "사용자: JWT, Agent: registration_token → agent_secret"],
        ["데이터 보안", "고객 코드/데이터 절대 클라우드 전송 금지"],
        ["비밀 관리", "agent_secret은 고객 서버만 보유 (클라우드는 해시만)"],
        ["라이센스", "24시간 주기 검증, 72시간 오프라인 grace period"],
    ]
    _add_table(doc, ["영역", "설명"], sec)

    # 7. 확장성
    doc.add_heading("7. 확장성 설계", level=1)
    scale = [
        ["Cloud 수평 확장", "FastAPI + Redis로 WebSocket Hub 수평 확장"],
        ["Agent 확장", "고객 서버 1대당 1 Agent (N개 프로젝트 관리)"],
        ["고가용성", "로드밸런서 + 다중 인스턴스 + Redis Pub/Sub 브로커링"],
        ["재연결", "지수 백오프 (1s → max 5분) + 밀린 상태 자동 보고"],
    ]
    _add_table(doc, ["항목", "설명"], scale)

    out = SPEC_DIR / f"24SevenClaw_시스템아키텍처_{TODAY}.docx"
    doc.save(str(out))
    return out


# ============================================================
# 5. 일일 작업 보고서
# ============================================================

def generate_daily_report() -> Path:
    """오늘의 Git 커밋/변경 기반 일일 작업 보고서 생성."""
    doc = Document()
    _style_doc(doc)
    _add_cover(doc, "일일 작업 보고서", f"{TODAY} 작업 내역")

    # Git 로그 수집
    doc.add_heading("1. 오늘의 커밋 내역", level=1)
    try:
        git_log = subprocess.check_output(
            ["git", "log", "--since=midnight", "--oneline", "--no-merges"],
            cwd=str(ROOT),
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
    except subprocess.CalledProcessError:
        git_log = "(git log 조회 실패)"

    if git_log:
        for line in git_log.split("\n"):
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("오늘 커밋 없음")

    # 변경 파일 통계
    doc.add_heading("2. 변경 파일 통계", level=1)
    try:
        diff_stat = subprocess.check_output(
            ["git", "diff", "--stat", "HEAD~5", "HEAD"],
            cwd=str(ROOT),
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
    except subprocess.CalledProcessError:
        diff_stat = "(diff stat 조회 실패)"

    if diff_stat:
        _add_code_block(doc, diff_stat)
    else:
        doc.add_paragraph("변경 사항 없음")

    # fix_plan 상태
    doc.add_heading("3. Fix Plan 상태", level=1)
    fix_plan = ROOT / ".ralph" / "fix_plan.md"
    if fix_plan.exists():
        content = fix_plan.read_text(encoding="utf-8")
        completed = content.count("- [x]")
        pending = content.count("- [ ]")
        blocked = content.count("- [!]")
        _add_table(doc, ["상태", "개수"], [
            ["완료", str(completed)],
            ["미완료", str(pending)],
            ["건너뜀", str(blocked)],
        ])
    else:
        doc.add_paragraph("fix_plan.md 파일 없음")

    out = SPEC_DIR / f"24SevenClaw_일일보고서_{TODAY}.docx"
    doc.save(str(out))
    return out


# ============================================================
# CLI
# ============================================================

def main() -> None:
    parser = argparse.ArgumentParser(description="24SevenClaw 사양 문서 생성기")
    parser.add_argument("--all", action="store_true", help="모든 문서 생성")
    parser.add_argument("--prd", action="store_true", help="기획서 생성")
    parser.add_argument("--api", action="store_true", help="API 정의서 생성")
    parser.add_argument("--tech", action="store_true", help="기술설계서 생성")
    parser.add_argument("--arch", action="store_true", help="시스템 아키텍처 생성")
    parser.add_argument("--daily", action="store_true", help="일일 보고서 생성")
    args = parser.parse_args()

    # 인자 없으면 --daily (일일 보고서만 자동 생성, 기타 문서는 명시적 요청 필요)
    if not any([args.all, args.prd, args.api, args.tech, args.arch, args.daily]):
        args.daily = True

    generators = []
    if args.all or args.prd:
        generators.append(("기획서 (PRD)", generate_prd))
    if args.all or args.api:
        generators.append(("API 정의서", generate_api_spec))
    if args.all or args.tech:
        generators.append(("기술설계서", generate_tech_design))
    if args.all or args.arch:
        generators.append(("시스템 아키텍처", generate_architecture))
    if args.all or args.daily:
        generators.append(("일일 보고서", generate_daily_report))

    for name, gen_fn in generators:
        print(f"  생성 중: {name}...", end=" ", flush=True)
        out_path = gen_fn()
        print(f"완료 → {out_path.relative_to(ROOT)}")

    print(f"\n  총 {len(generators)}개 문서 생성 완료 ({SPEC_DIR.relative_to(ROOT)}/)")


if __name__ == "__main__":
    main()
